from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Protocol

from docx import Document

from src.memory import save_tailoring_run
from src.utils import clean_html_to_text


class TextGenerator(Protocol):
    model: str

    def generate(self, prompt: str, system: str = "") -> str:
        ...


@dataclass(frozen=True)
class TailoringResult:
    run_id: int
    analysis_text: str
    tailored_resume_markdown: str
    markdown_path: str
    docx_path: str
    model: str


def tailor_resume_for_job(
    db_path: str | Path,
    job: dict[str, Any],
    active_resume: dict[str, Any],
    resume_text: str,
    user_notes: str,
    generator: TextGenerator,
    config: dict[str, Any] | None = None,
) -> TailoringResult:
    config = config or {}
    job_id = int(job["id"])
    resume_version_id = int(active_resume["id"])
    prompt = _build_prompt(job, resume_text, user_notes, config)
    output = generator.generate(prompt, system=_system_prompt())
    analysis, resume_markdown = _split_output(output)
    analysis = _clean_model_output(analysis)
    resume_markdown = _clean_model_output(resume_markdown)

    output_dir = Path(str(config.get("output_dir") or "data/profile/tailored_resumes"))
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    safe_title = _safe_filename(str(job.get("title") or "job"))
    run_dir = output_dir / f"job_{job_id}" / timestamp
    run_dir.mkdir(parents=True, exist_ok=True)
    markdown_path = run_dir / f"{safe_title}_tailored_resume.md"
    docx_path = run_dir / f"{safe_title}_tailored_resume.docx"

    markdown_path.write_text(resume_markdown, encoding="utf-8")
    _write_docx(resume_markdown, docx_path)
    run_id = save_tailoring_run(
        db_path,
        job_id,
        resume_version_id,
        generator.model,
        user_notes,
        analysis,
        markdown_path.as_posix(),
        docx_path.as_posix(),
    )
    return TailoringResult(
        run_id=run_id,
        analysis_text=analysis,
        tailored_resume_markdown=resume_markdown,
        markdown_path=markdown_path.as_posix(),
        docx_path=docx_path.as_posix(),
        model=generator.model,
    )


def _system_prompt() -> str:
    return (
        "You are Hermes, a local-first resume tailoring assistant. "
        "You help tailor an existing resume to a job description without inventing experience. "
        "Be truthful, concrete, and conservative. Never claim tools, platforms, employers, titles, "
        "certifications, degrees, or outcomes unless they appear in the source resume. "
        "You may rephrase, reorder, emphasize, and summarize existing evidence."
    )


def _build_prompt(
    job: dict[str, Any],
    resume_text: str,
    user_notes: str,
    config: dict[str, Any],
) -> str:
    max_resume_chars = int(config.get("max_resume_chars") or 12000)
    max_job_chars = int(config.get("max_job_chars") or 16000)
    job_description = clean_html_to_text(str(job.get("description_text") or ""))[:max_job_chars]
    resume_excerpt = resume_text[:max_resume_chars]
    title = str(job.get("title") or "")
    company = str(job.get("company") or "")
    location = str(job.get("location") or "")
    notes = user_notes.strip() or "No additional user notes."
    return f"""
Return exactly two sections with these headings:

## Fit Analysis
Use this exact structure:

### Fit Verdict
One short paragraph explaining the overall fit.

### Evidence Map
Create a Markdown table with these columns:
JD requirement | Resume evidence | How to position it | Confidence

Use concrete evidence from the source resume. If the evidence is adjacent rather than direct, say so.

### Gaps / Do Not Overclaim
List missing or weak areas. Clearly state what must not be claimed.

### Recommended Emphasis
List the themes that should be emphasized in the tailored resume.

## Tailored Resume
Create a concise, truthful tailored resume draft in Markdown.
Keep it ATS-friendly.
Do not wrap the resume in ```markdown or any code fence.
Do not include horizontal rules such as ---.
Use only evidence from the source resume.
Do not invent employer names, dates, degrees, certifications, metrics, tools, or achievements.
If the job asks for something missing from the resume, do not add it as experience; place it in a brief "Relevant exposure / adjacent strengths" style only if supported.

Job:
Title: {title}
Company: {company}
Location: {location}

User guidance:
{notes}

Job description:
{job_description}

Source resume:
{resume_excerpt}
""".strip()


def _split_output(output: str) -> tuple[str, str]:
    output = _clean_model_output(output)
    match = re.search(r"^## Tailored Resume\s*$", output, flags=re.MULTILINE)
    if not match:
        return output.strip(), output.strip()
    analysis = output[: match.start()].strip()
    resume = output[match.start() :].strip()
    return analysis, resume


def _clean_model_output(value: str) -> str:
    text = value.strip()
    text = re.sub(r"^```(?:markdown|md)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```\s*$", "", text)
    cleaned_lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip() in {"---", "***", "___"}:
            continue
        if line.strip().casefold() in {"`markdown", "markdown"}:
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def _write_docx(markdown_text: str, path: Path) -> None:
    document = Document()
    for raw_line in _clean_model_output(markdown_text).splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("•"):
            bullet = line[2:].strip() if line.startswith("- ") else line.lstrip("•").strip()
            document.add_paragraph(_strip_markdown(bullet), style="List Bullet")
        elif _looks_like_markdown_table(line):
            continue
        else:
            document.add_paragraph(_strip_markdown(line))
    document.save(path)


def _strip_markdown(value: str) -> str:
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
    value = re.sub(r"\*(.*?)\*", r"\1", value)
    value = re.sub(r"`([^`]*)`", r"\1", value)
    return value


def _looks_like_markdown_table(value: str) -> bool:
    stripped = value.strip()
    if not stripped.startswith("|"):
        return False
    return stripped.endswith("|")


def _safe_filename(value: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", value).strip("_")
    return safe[:80] or "tailored_resume"
