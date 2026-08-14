from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from src.memory import save_resume_version
from src.scorer import ScoreResult

SUPPORTED_RESUME_EXTENSIONS = {".docx", ".pdf", ".txt", ".md"}
RESUME_TERMS = [
    "airflow",
    "analytics",
    "aws",
    "banking",
    "business intelligence",
    "customer analytics",
    "data analysis",
    "data engineering",
    "data science",
    "excel",
    "financial services",
    "fintech",
    "forecasting",
    "machine learning",
    "power bi",
    "python",
    "risk analytics",
    "risk modeling",
    "snowflake",
    "spark",
    "sql",
    "statistics",
    "tableau",
]


@dataclass(frozen=True)
class ResumeFit:
    score: int
    matched_terms: list[str]
    missing_terms: list[str]


@dataclass(frozen=True)
class ResumeScoreAdjustment:
    score_delta: int
    reasons: list[str]
    fit: ResumeFit | None


def save_uploaded_resume(
    db_path: str | Path,
    original_filename: str,
    content: bytes,
    profile_dir: str | Path = "data/profile",
) -> dict[str, Any]:
    extension = Path(original_filename).suffix.lower()
    if extension not in SUPPORTED_RESUME_EXTENSIONS:
        raise ValueError(f"Unsupported resume file type: {extension}")

    file_hash = hashlib.sha256(content).hexdigest()
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    safe_name = _safe_filename(Path(original_filename).stem)
    profile_path = Path(profile_dir)
    original_dir = profile_path / "resumes" / "original"
    extracted_dir = profile_path / "resumes" / "extracted"
    original_dir.mkdir(parents=True, exist_ok=True)
    extracted_dir.mkdir(parents=True, exist_ok=True)

    stored_path = original_dir / f"{timestamp}_{safe_name}{extension}"
    extracted_path = extracted_dir / f"{timestamp}_{safe_name}.txt"
    stored_path.write_bytes(content)

    extracted_text = extract_resume_text(original_filename, content)
    extracted_path.write_text(extracted_text, encoding="utf-8")
    resume_id = save_resume_version(
        db_path,
        original_filename,
        stored_path.as_posix(),
        extracted_path.as_posix(),
        file_hash,
    )
    return {
        "id": resume_id,
        "original_filename": original_filename,
        "stored_path": stored_path.as_posix(),
        "extracted_text_path": extracted_path.as_posix(),
        "file_hash": file_hash,
    }


def extract_resume_text(original_filename: str, content: bytes) -> str:
    extension = Path(original_filename).suffix.lower()
    if extension == ".docx":
        return _extract_docx_text(content)
    if extension == ".pdf":
        return _extract_pdf_text(content)
    if extension in {".txt", ".md"}:
        return _clean_text(content.decode("utf-8", errors="replace"))
    raise ValueError(f"Unsupported resume file type: {extension}")


def load_resume_text(resume: dict[str, Any] | None) -> str:
    if not resume:
        return ""
    path = Path(str(resume.get("extracted_text_path") or ""))
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def score_resume_fit(job: dict[str, Any], resume_text: str) -> ResumeFit | None:
    if not resume_text.strip():
        return None
    resume_terms = _matched_terms(resume_text, RESUME_TERMS)
    if not resume_terms:
        return ResumeFit(score=0, matched_terms=[], missing_terms=[])

    job_text = " ".join(
        str(job.get(key) or "")
        for key in ("title", "description_text", "matched_skills", "positive_reasons")
    )
    job_terms = _matched_terms(job_text, RESUME_TERMS)
    matched = sorted(set(resume_terms) & set(job_terms))
    missing = sorted(set(job_terms) - set(resume_terms))
    score = round(100 * len(matched) / max(1, len(set(job_terms)))) if job_terms else 0
    return ResumeFit(score=score, matched_terms=matched, missing_terms=missing[:8])


def adjust_score_for_resume(
    result: ScoreResult,
    job: dict[str, Any],
    resume_text: str,
    config: dict[str, Any] | None = None,
) -> ScoreResult:
    adjustment = calculate_resume_score_adjustment(job, resume_text, config)
    if not adjustment.fit or adjustment.score_delta == 0:
        return result

    positive_reasons = list(result.positive_reasons)
    negative_reasons = list(result.negative_reasons)
    if adjustment.score_delta > 0:
        positive_reasons.extend(adjustment.reasons)
    else:
        negative_reasons.extend(adjustment.reasons)

    score = max(0, min(100, result.score + adjustment.score_delta))
    return ScoreResult(
        score=score,
        matched_skills=result.matched_skills,
        missing_skills=result.missing_skills,
        positive_reasons=positive_reasons,
        negative_reasons=negative_reasons,
    )


def apply_resume_adjustment_to_job(
    job: dict[str, Any],
    resume_text: str,
    config: dict[str, Any] | None = None,
) -> dict[str, Any]:
    adjustment = calculate_resume_score_adjustment(job, resume_text, config)
    adjusted = dict(job)
    adjusted["resume_fit_score"] = adjustment.fit.score if adjustment.fit else None
    adjusted["resume_adjustment"] = adjustment.score_delta
    adjusted["resume_reasons"] = adjustment.reasons
    if adjustment.score_delta:
        adjusted["base_score_before_resume"] = int(adjusted.get("score") or 0)
        adjusted["score"] = max(0, min(100, int(adjusted.get("score") or 0) + adjustment.score_delta))
    return adjusted


def calculate_resume_score_adjustment(
    job: dict[str, Any],
    resume_text: str,
    config: dict[str, Any] | None = None,
) -> ResumeScoreAdjustment:
    config = config or {}
    if not config.get("enabled", False):
        return ResumeScoreAdjustment(score_delta=0, reasons=[], fit=None)

    fit = score_resume_fit(job, resume_text)
    if fit is None:
        return ResumeScoreAdjustment(score_delta=0, reasons=[], fit=None)

    high_threshold = int(config.get("high_fit_threshold", 70))
    low_threshold = int(config.get("low_fit_threshold", 35))
    high_bonus = int(config.get("high_fit_bonus", 8))
    moderate_bonus = int(config.get("moderate_fit_bonus", 4))
    low_penalty = int(config.get("low_fit_penalty", 6))

    if fit.score >= high_threshold:
        terms = ", ".join(fit.matched_terms[:6]) or "resume terms"
        return ResumeScoreAdjustment(
            score_delta=high_bonus,
            reasons=[f"Resume fit is strong ({fit.score}/100): {terms}"],
            fit=fit,
        )
    if fit.score >= low_threshold:
        terms = ", ".join(fit.matched_terms[:6]) or "resume terms"
        return ResumeScoreAdjustment(
            score_delta=moderate_bonus,
            reasons=[f"Resume fit is moderate ({fit.score}/100): {terms}"],
            fit=fit,
        )
    missing = ", ".join(fit.missing_terms[:6]) or "job keywords"
    return ResumeScoreAdjustment(
        score_delta=-low_penalty,
        reasons=[f"Resume fit is weak ({fit.score}/100); missing from resume: {missing}"],
        fit=fit,
    )


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_text.append(cell.text)
    return _clean_text("\n".join(paragraphs + table_text))


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    parts = [page.extract_text() or "" for page in reader.pages]
    return _clean_text("\n".join(parts))


def _matched_terms(text: str, terms: list[str]) -> list[str]:
    normalized = _normalize(text)
    return [term for term in terms if _normalize(term) in normalized]


def _clean_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"\n{3,}", "\n\n", value)
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def _normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value.casefold()).strip()


def _safe_filename(value: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", value).strip("_")
    return safe or "resume"
